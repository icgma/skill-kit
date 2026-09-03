#!/usr/bin/env python3
"""OPC package plumbing shared by every tool: relationship-target
resolution, safe extraction, atomic writes.

Three failure classes this module exists to kill:
- each tool hand-rolling its own Target join drifts at the edges
  (`../`, percent-encoding, TargetMode="External", package-root
  absolute paths) -- one resolver, one set of edges;
- extracting a foreign docx without a path guard lets a hostile entry
  (`../x`, an absolute path, a symlink) write outside the destination;
- writing onto the target in place (tree.write / zip rebuild) leaves a
  half-written file when interrupted. The atomic pattern -- temp file
  in the SAME directory, then os.replace -- either fully succeeds or
  changes nothing. Same directory matters: /tmp is often a different
  mount, and a cross-device "move" degrades to copy+delete, which is
  exactly the non-atomic write this exists to prevent.
"""
from __future__ import annotations

import os
import re
import tempfile
import zipfile
from pathlib import Path, PurePosixPath
from urllib.parse import unquote
import sys

# Windows consoles/pipes default to the legacy ANSI code page (cp1252/cp936):
# the engine's CJK messages must never crash the gate there.
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")


# ---------------------------------------------------------------- rels

def rel_target(*, target: str | None, base_part: str,
               target_mode: str | None = None) -> str | None:
    """Resolve a relationship Target against the part declaring it.
    KEYWORD-ONLY (review #14): a positional call with the arguments
    reversed resolved every header relationship to the base part
    itself, and a downstream fallback masked it -- the type system
    cannot catch a str/str swap, so the signature must.

    -> normalized part name ("word/media/image1.png"),
       None for External / empty targets,
       ValueError for targets that escape the package root (malformed;
       every caller must notice, not silently mis-resolve).

    base_part is the OWNING part ("word/document.xml" for
    word/_rels/document.xml.rels); "" means the package root
    (_rels/.rels). Percent-encoding is decoded PER SEGMENT and a decoded
    segment containing a separator is rejected -- decoding first and
    splitting second would let "a%2F..%2F.." smuggle a `..` past the
    stack. A leading "/" is package-root-absolute, not host filesystem.
    """
    if not target or target_mode == "External":
        return None
    stack = ([] if target.startswith("/")
             else list(PurePosixPath(base_part).parent.parts))
    if stack == ["."]:
        stack = []
    for seg in target.lstrip("/").split("/"):
        seg = unquote(seg)
        if "/" in seg or "\\" in seg:
            raise ValueError(f"encoded separator in Target segment: "
                             f"{target!r}")
        if seg in ("", "."):
            continue
        if seg == "..":
            if not stack:
                raise ValueError(f"Target escapes package root: {target!r}")
            stack.pop()
        else:
            stack.append(seg)
    return "/".join(stack) or None


def rels_owner(rels_name: str) -> str:
    """dir/_rels/name.rels -> the part owning it (dir/name);
    _rels/.rels -> "" (the package root)."""
    p = PurePosixPath(rels_name)
    owner = p.parent.parent / p.name[:-len(".rels")]
    s = str(owner)
    return "" if s in (".rels", "/.rels") else s.lstrip("/")


# ------------------------------------------------- part resolution (OPC)

_RT = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/"
_PR = "{http://schemas.openxmlformats.org/package/2006/relationships}"
#: relationship types whose targets carry TEXT (a story). Anything here
#: is walked by read/validate/revisions; anything not is data.
STORY_RT = {_RT + k for k in ("header", "footer", "footnotes", "endnotes",
                              "comments")}


def _xml(reader, name):
    try:
        data = reader(name)
    except Exception:
        return None
    if data is None:
        return None
    try:
        from lxml import etree
        return etree.fromstring(data)
    except Exception:
        return None


def main_part(reader, names) -> str:
    """The main document part, resolved through _rels/.rels. NEVER a
    hardcoded path: OPC lets a package name it anything, and Word opens
    such a package fine.

    Shared by read/validate/revisions/track ON PURPOSE. Each tool used to
    decide for itself -- read resolved via rels, revisions matched the
    `word/` prefix -- so a main part at `doc/main.xml` was rendered by the
    reader, PASSED by the validator, and silently written through
    unchanged by `revisions.py --accept` with exit code 0 and the message
    "processed 0 revisions" (dxv2-6 review B4). Half-support is worse
    than no support: it reports success."""
    names = set(names)
    rr = _xml(reader, "_rels/.rels")
    if rr is not None:
        for rel in rr.findall(_PR + "Relationship"):
            if (rel.get("Type") or "") == _RT + "officeDocument":
                try:
                    nm = rel_target(target=rel.get("Target"), base_part="")
                except ValueError:
                    continue
                if nm and nm in names:
                    return nm
    return "word/document.xml"


def parts_by_type(reader, names) -> dict:
    """{"comments": "word/cmts9.xml", "header": [...], ...} -- EVERY
    consumer's answer to "where does this story live", derived once from
    the main part's relationships.

    This exists because `story_parts` alone was not enough: read.py kept
    THREE separate hand-rolled walkers (comments/notes, headers/footers,
    and the styles/numbering pair), each opening
    `word/_rels/document.xml.rels` and resolving against
    `base_part="word/document.xml"`. Adding a shared resolver without
    deleting them just made a fourth implementation, so a package whose
    main part is `doc/main.xml` rendered its body while its comment and
    footnote bodies vanished -- and validate still said PASSED (dxv2-7
    review P1.1). One concept, one implementation."""
    names = set(names)
    main = main_part(reader, names)
    out: dict = {"document": main}
    if main not in names:
        return out
    rels = f"{PurePosixPath(main).parent}/_rels/" \
           f"{PurePosixPath(main).name}.rels"
    rels = rels.lstrip("./")
    dr = _xml(reader, rels)
    KINDS = ("comments", "footnotes", "endnotes", "styles", "numbering",
             "settings", "commentsExtended", "commentsIds")
    MULTI = ("header", "footer")
    for k in MULTI:
        out[k] = []
    if dr is not None:
        for rel in dr.findall(_PR + "Relationship"):
            ty = (rel.get("Type") or "")
            try:
                nm = rel_target(target=rel.get("Target"), base_part=main)
            except ValueError:
                continue
            if not nm or nm not in names:
                continue
            for k in KINDS:
                if ty == _RT + k:
                    out.setdefault(k, nm)
            for k in MULTI:
                if ty == _RT + k:
                    out[k].append(nm)
    # name-pattern fallback, scoped to the main part's own directory so a
    # package cannot pick up a stray word/comments.xml that nothing links
    base = str(PurePosixPath(main).parent)
    for k in KINDS:
        cand = f"{base}/{k}.xml".lstrip("/")
        if k not in out and cand in names:
            out[k] = cand
    for k in MULTI:
        out[k] = sorted(set(out[k]) | {
            n for n in names
            if re.fullmatch(rf"{re.escape(base)}/{k}\d*\.xml", n)})
    return out


def story_parts(reader, names) -> set:
    """Every text-bearing part: the main document plus whatever its rels
    declare as header/footer/notes/comments -- by TYPE, not by name."""
    by = parts_by_type(reader, names)
    out = {by["document"]} if by.get("document") in set(names) else set()
    for k in ("comments", "footnotes", "endnotes"):
        if by.get(k):
            out.add(by[k])
    for k in ("header", "footer"):
        out |= set(by.get(k) or ())
    if not out:
        out = {n for n in names
               if re.fullmatch(r"word/(header\d*|footer\d*|footnotes|"
                               r"endnotes|comments)\.xml", n)}
    return out


# ---------------------------------------------------------------- unpack

#: Decompression caps. Generous for any real document (the largest
#: corpus file is ~21 MB packed), tight enough that a zip bomb -- a few
#: KB expanding to tens of GB -- dies with a clear message instead of
#: exhausting memory/disk. Declared sizes in the central directory can
#: lie, so enforcement is on the actually-decompressed bytes.
_MAX_PART = 256 * 1024 * 1024        # one entry
_MAX_TOTAL = 1024 * 1024 * 1024      # whole package


class CappedZip:
    """ZipFile wrapper enforcing the SAME decompression caps as
    unpack() on every read. The caps lived only in unpack(); prep/
    read/validate/revisions called ZipFile.read() directly, so a zip
    bomb entered through the main entry points unbounded
    (Ultra-review P1). Business scripts use THIS, never raw
    ZipFile.read()."""

    def __init__(self, src):
        self._z = zipfile.ZipFile(src)
        self._seen: dict = {}        # name -> decompressed size (once)
        names = [i.filename for i in self._z.infolist() if not i.is_dir()]
        from collections import Counter as _C
        #: EXPOSED, not enforced: validate must be able to open an
        #: ambiguous package in order to REPORT it (refusing at the door
        #: turned its own diagnostic into a crash). Paths that have to
        #: PICK one (unpack, the view) refuse; the reporter does not.
        self.duplicates = sorted(k for k, v in _C(names).items() if v > 1)

    def namelist(self):
        return self._z.namelist()

    def infolist(self):
        return self._z.infolist()

    def read(self, name) -> bytes:
        # streamed with the cap enforced per chunk: ZipFile.read()
        # would materialize the whole bomb before we could refuse.
        # The package total counts each DISTINCT entry once -- re-reading
        # one entry must not re-accumulate (a legit package copied part
        # by part was wrongly killed; dxv2-3 review P1.7).
        out, got = [], 0
        with self._z.open(name) as fh:
            while True:
                chunk = fh.read(1 << 20)
                if not chunk:
                    break
                got += len(chunk)
                if got > _MAX_PART:
                    raise SystemExit(
                        f"zip entry {name!r} decompresses past "
                        f"{_MAX_PART} bytes: refusing (zip-bomb defense)")
                prior = self._seen.get(name, 0)
                if sum(self._seen.values()) - prior + got > _MAX_TOTAL:
                    raise SystemExit(
                        f"package decompresses past {_MAX_TOTAL} bytes "
                        "total: refusing (zip-bomb defense)")
                out.append(chunk)
        self._seen[name] = got
        return b"".join(out)

    def __enter__(self):
        return self

    def __exit__(self, *a):
        self._z.close()
        return False


def unpack(src, dest) -> int:
    """Extract a .docx into dest with hostile-entry defense: entries
    that escape dest (`../x`, absolute paths, same-prefix siblings via
    resolve()) are skipped; symlink entries land as PLAIN FILES holding
    the link text (we write bytes ourselves and never recreate the
    link -- a symlink inside the tree would turn every later write into
    a write-through-attacker-path); decompressed sizes are capped
    (zip-bomb defense). -> number of files written."""
    dest = Path(dest)
    dest.mkdir(parents=True, exist_ok=True)
    base = dest.resolve()
    n, total = 0, 0
    with zipfile.ZipFile(src) as z:
        # duplicate entry names: OPC forbids them and every reader picks
        # a different one. Unpacking silently kept the LAST, so the view
        # (and every edit) ran against a version nobody chose
        # (dxv2-4 review P1.9). Refuse at the door -- the package is
        # ambiguous, not repairable by guessing.
        seen: dict = {}
        for info in z.infolist():
            if info.is_dir():
                continue
            seen[info.filename] = seen.get(info.filename, 0) + 1
        dups = sorted(k for k, v in seen.items() if v > 1)
        if dups:
            raise SystemExit(
                "package contains duplicate entries: "
                + ", ".join(dups[:5])
                + (f" (+{len(dups) - 5} more)" if len(dups) > 5 else "")
                + " -- OPC requires unique part names and readers "
                "disagree on which one wins; fix the package before "
                "editing (unzip -l to inspect)")
        for info in z.infolist():
            tgt = (base / info.filename.lstrip("/")).resolve()
            if not (tgt == base or base in tgt.parents):
                continue
            if info.is_dir():
                tgt.mkdir(parents=True, exist_ok=True)
                continue
            with z.open(info) as fh:
                data = fh.read(_MAX_PART + 1)
            if len(data) > _MAX_PART:
                raise SystemExit(
                    f"{info.filename} decompresses past "
                    f"{_MAX_PART >> 20} MiB -- refusing (zip bomb?)")
            total += len(data)
            if total > _MAX_TOTAL:
                raise SystemExit(
                    f"package decompresses past {_MAX_TOTAL >> 20} MiB "
                    "in total -- refusing (zip bomb?)")
            tgt.parent.mkdir(parents=True, exist_ok=True)
            tgt.write_bytes(data)
            n += 1
    return n


# ---------------------------------------------------------------- atomic

def atomic_write(path, data: bytes) -> None:
    """Write bytes so `path` is either fully old or fully new; a failed
    attempt leaves no temp file behind."""
    path = Path(path)
    fd, tmp = tempfile.mkstemp(dir=str(path.parent),
                               prefix=path.name + ".", suffix=".tmp")
    try:
        with os.fdopen(fd, "wb") as f:
            f.write(data)
        os.replace(tmp, path)
    except BaseException:
        try:
            os.unlink(tmp)
        except OSError:
            pass
        raise


def atomic_write_tree(tree, path) -> None:
    """lxml tree/element -> atomic file, in the serialization every tool
    here uses (declaration + UTF-8 + standalone)."""
    from lxml import etree
    atomic_write(path, etree.tostring(
        tree, xml_declaration=True, encoding="UTF-8", standalone=True))


def atomic_zip_rewrite(src, dst, replace: dict[str, bytes] | None = None,
                       ) -> None:
    """Copy zip src -> dst, substituting entries named in `replace`
    (name -> new bytes). Atomic: assembled next to dst, then replaced.
    With no replacements the bytes are copied verbatim (no recompression
    churn)."""
    src, dst = Path(src), Path(dst)
    with zipfile.ZipFile(src) as _z:
        _n = [i.filename for i in _z.infolist() if not i.is_dir()]
    if len(_n) != len(set(_n)):
        from collections import Counter as _C
        d = sorted(k for k, v in _C(_n).items() if v > 1)
        # THE write primitive refuses too, so no editor path can copy an
        # ambiguous package forward (dxv2-5 review P1.4)
        raise SystemExit(
            "refusing to rewrite a package with duplicate entries: "
            + ", ".join(d[:5]))
    if not replace:
        atomic_write(dst, src.read_bytes())
        return
    fd, tmp = tempfile.mkstemp(dir=str(dst.parent),
                               prefix=dst.name + ".", suffix=".tmp")
    try:
        with os.fdopen(fd, "wb") as f, \
                CappedZip(src) as zin, \
                zipfile.ZipFile(f, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                # capped read, and LAZY: an entry being replaced is
                # never decompressed at all (the eager default argument
                # read it first -- dxv2-4 review P1.9)
                if item.filename in replace:
                    zout.writestr(item, replace[item.filename])
                else:
                    zout.writestr(item, zin.read(item.filename))
        os.replace(tmp, dst)
    except BaseException:
        try:
            os.unlink(tmp)
        except OSError:
            pass
        raise


def _main() -> int:
    """CLI so the repack step gets the same atomicity as the tools:
        python scripts/opc.py pack work/unpacked out.docx
        python scripts/opc.py unpack in.docx work/unpacked
    A shell `zip -Xr` writes the target in place -- interrupt it and the
    docx is half a zip."""
    import sys as _sys
    a = _sys.argv[1:]
    if len(a) == 3 and a[0] == "pack":
        atomic_zip_dir(a[1], a[2])
        suf = Path(a[2]).suffix.lower()
        note = (f"  (main part typed as {suf[1:]})"
                if suf in MAIN_CT else "")
        print(f"packed {a[1]} -> {a[2]}{note}")
        return 0
    if len(a) == 3 and a[0] == "unpack":
        n = unpack(a[1], a[2])
        print(f"unpacked {n} files -> {a[2]}")
        return 0
    print(__doc__ and _main.__doc__)
    return 2


def file_hash(path) -> str:
    """Content fingerprint of a packed file."""
    import hashlib
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(1 << 16), b""):
            h.update(chunk)
    return h.hexdigest()


def tree_hash(root) -> str:
    """Content fingerprint of an unpacked directory: every file's path
    and bytes, order-independent of filesystem quirks. mtimes are
    useless here -- unzip sets them arbitrarily."""
    import hashlib
    root = Path(root)
    h = hashlib.sha256()
    for p in sorted(root.rglob("*")):
        if p.is_file():
            h.update(str(p.relative_to(root)).encode())
            h.update(p.read_bytes())
    return h.hexdigest()



#: The main part's content type is what tells Word "document" from
#: "template". prep.py retypes a .dotx working copy to document so
#: python-docx can open it; packing back out to a .dotx has to undo
#: that, or the deliverable is a DOCX wearing a .dotx name and Word
#: will not offer it as a template (dxv2-7 review P1.6).
MAIN_CT = {
    ".docx": "application/vnd.openxmlformats-officedocument."
             "wordprocessingml.document.main+xml",
    ".dotx": "application/vnd.openxmlformats-officedocument."
             "wordprocessingml.template.main+xml",
    ".docm": "application/vnd.ms-word.document.macroEnabled.main+xml",
    ".dotm": "application/vnd.ms-word.template.macroEnabled.main+xml",
}


def retype_main(ct_xml: bytes, suffix: str) -> bytes:
    """Rewrite [Content_Types].xml's MAIN-part override to match `suffix`.
    Returns the input unchanged when the suffix is unknown or already
    right -- never invents an override that was not there."""
    want = MAIN_CT.get(suffix.lower())
    if not want:
        return ct_xml
    out = ct_xml
    for ct in MAIN_CT.values():
        if ct == want:
            continue
        if ct.encode() in out:
            out = out.replace(ct.encode(), want.encode())
    return out


def atomic_zip_dir(src_dir, dst) -> None:
    """Zip a directory into dst atomically (entry names are /-separated
    paths relative to src_dir). The main part's content type is retyped
    to match dst's extension."""
    src_dir, dst = Path(src_dir), Path(dst)
    fd, tmp = tempfile.mkstemp(dir=str(dst.parent),
                               prefix=dst.name + ".", suffix=".tmp")
    try:
        with os.fdopen(fd, "wb") as f, \
                zipfile.ZipFile(f, "w", zipfile.ZIP_DEFLATED) as zout:
            for p in sorted(src_dir.rglob("*")):
                if not p.is_file():
                    continue
                name = str(p.relative_to(src_dir)).replace("\\", "/")
                if name == "[Content_Types].xml":
                    zout.writestr(name, retype_main(p.read_bytes(),
                                                    dst.suffix))
                else:
                    zout.write(p, name)
        os.replace(tmp, dst)
    except BaseException:
        try:
            os.unlink(tmp)
        except OSError:
            pass
        raise


def err(code: str, what: str, where: str = "", fix: str = "") -> SystemExit:
    """The ONE tool-error format: `E_CODE: what | where | try: fix`.
    English, dense, greppable, locating. Tests pin the CODE; the prose
    can be reworded freely. Guidance lines (non-error stdout) stay in
    the docs' language -- two channels, two voices, so a model can tell
    'this is directions' from 'this broke' at a glance."""
    parts = [f"{code}: {what}"]
    if where:
        parts.append(where)
    if fix:
        parts.append(f"try: {fix}")
    return SystemExit(" | ".join(parts))


class edit_docx:
    """python-docx over the unpacked working copy -- the bridge that
    makes work/unpacked/ the ONLY editable surface:

        with opc.edit_docx("work/unpacked") as doc:
            doc.paragraphs[0].add_run(" ...")

    Enter: zip the tree to a temp .docx and open it with python-docx.
    Exit (no exception): save + unpack back over the tree atomically-ish
    (unpack to a sibling temp dir, then swap). Exception inside the
    block: nothing is written. python-docx re-serializes the parts it
    touched -- byte shape may change, semantics do not; the gate is
    semantic."""

    def __init__(self, unpacked):
        self.dir = Path(unpacked)
        self.tmp = None
        self.doc = None

    def __enter__(self):
        import tempfile
        import docx
        self.tmp = Path(tempfile.mkdtemp(prefix="edx"))
        self.pkg = self.tmp / "pkg.docx"
        atomic_zip_dir(self.dir, self.pkg)
        self.doc = docx.Document(str(self.pkg))
        return self.doc

    def __exit__(self, et, ev, tb):
        import shutil
        try:
            if et is None:
                self.doc.save(str(self.pkg))
                # stage the new tree BESIDE the target (same fs), so the
                # swap is two metadata renames -- the earlier version
                # cross-device-copied from /tmp between the renames,
                # leaving a whole-tree-copy-sized crash window with the
                # workspace absent (review item)
                new = self.dir.parent / (self.dir.name + ".new~")
                old = self.dir.parent / (self.dir.name + ".old~")
                for leftover in (new, old):
                    if leftover.exists():
                        shutil.rmtree(leftover)
                unpack(self.pkg, new)
                self.dir.rename(old)
                try:
                    new.rename(self.dir)
                except BaseException:
                    old.rename(self.dir)      # roll back the swap
                    raise
                shutil.rmtree(old, ignore_errors=True)
        finally:
            shutil.rmtree(self.tmp, ignore_errors=True)
        return False


if __name__ == "__main__":
    import sys
    sys.exit(_main())
